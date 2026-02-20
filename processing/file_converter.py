"""File conversion coordinator - converts various formats to text or images for Claude."""

import logging
from dataclasses import dataclass, field
from pathlib import Path

from processing.image_handler import SUPPORTED_IMAGE_EXTENSIONS, preprocess_image
from processing.pdf_handler import extract_text_from_pdf, pdf_to_images

logger = logging.getLogger(__name__)


@dataclass
class ConversionResult:
    """Result of file conversion."""
    original_path: Path
    text_content: str | None = None
    image_paths: list[Path] = field(default_factory=list)
    conversion_type: str = "unknown"  # text, image, mixed
    error: str | None = None


def convert_docx(file_path: Path) -> ConversionResult:
    """Convert DOCX to text, including table content."""
    from docx import Document

    doc = Document(str(file_path))
    parts = []

    for element in doc.element.body:
        tag = element.tag.split("}")[-1] if "}" in element.tag else element.tag
        if tag == "p":
            # Find the corresponding paragraph
            for para in doc.paragraphs:
                if para._element is element:
                    if para.text.strip():
                        parts.append(para.text)
                    break
        elif tag == "tbl":
            # Find the corresponding table
            for table in doc.tables:
                if table._element is element:
                    rows = []
                    for row in table.rows:
                        cells = [cell.text.strip() for cell in row.cells]
                        rows.append(" | ".join(cells))
                    if rows:
                        # Markdown-style table
                        parts.append(rows[0])
                        parts.append(" | ".join(["---"] * len(table.rows[0].cells)))
                        parts.extend(rows[1:])
                    break

    text = "\n".join(parts)
    return ConversionResult(
        original_path=file_path,
        text_content=text if text.strip() else None,
        conversion_type="text",
    )


def convert_xlsx(file_path: Path) -> ConversionResult:
    """Convert Excel to Markdown tables."""
    from openpyxl import load_workbook

    wb = load_workbook(str(file_path), data_only=True)
    parts = []

    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        parts.append(f"## Sheet: {sheet_name}\n")

        rows = list(ws.iter_rows(values_only=True))
        if not rows:
            continue

        # Filter empty rows
        rows = [r for r in rows if any(c is not None for c in r)]
        if not rows:
            continue

        # Header
        header = [str(c) if c is not None else "" for c in rows[0]]
        parts.append("| " + " | ".join(header) + " |")
        parts.append("| " + " | ".join(["---"] * len(header)) + " |")

        # Data rows
        for row in rows[1:]:
            cells = [str(c) if c is not None else "" for c in row]
            # Pad or truncate to match header length
            while len(cells) < len(header):
                cells.append("")
            parts.append("| " + " | ".join(cells[:len(header)]) + " |")

        parts.append("")

    wb.close()
    text = "\n".join(parts)
    return ConversionResult(
        original_path=file_path,
        text_content=text if text.strip() else None,
        conversion_type="text",
    )


def convert_file(file_path: str | Path, scan_dpi: int = 300, max_image_dim: int = 4096) -> ConversionResult:
    """Convert any supported file to text and/or images for Claude processing."""
    file_path = Path(file_path)
    suffix = file_path.suffix.lower()

    try:
        if suffix == ".pdf":
            text = extract_text_from_pdf(file_path)
            if text:
                logger.info("PDF has extractable text: %s", file_path.name)
                return ConversionResult(
                    original_path=file_path,
                    text_content=text,
                    conversion_type="text",
                )
            else:
                logger.info("PDF appears scanned, converting to images: %s", file_path.name)
                images = pdf_to_images(file_path, dpi=scan_dpi)
                processed = [preprocess_image(img, max_dimension=max_image_dim) for img in images]
                return ConversionResult(
                    original_path=file_path,
                    image_paths=processed,
                    conversion_type="image",
                )

        elif suffix == ".docx":
            return convert_docx(file_path)

        elif suffix == ".xlsx":
            return convert_xlsx(file_path)

        elif suffix in SUPPORTED_IMAGE_EXTENSIONS:
            processed = preprocess_image(file_path, max_dimension=max_image_dim)
            return ConversionResult(
                original_path=file_path,
                image_paths=[processed],
                conversion_type="image",
            )

        else:
            return ConversionResult(
                original_path=file_path,
                error=f"Unsupported file type: {suffix}",
            )

    except Exception as e:
        logger.error("Error converting %s: %s", file_path.name, e)
        return ConversionResult(
            original_path=file_path,
            error=str(e),
        )
